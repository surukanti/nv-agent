"""Document ingestion — load files, chunk, and index them."""

import logging
import os
import warnings
from pathlib import Path

from config import config
from kb.chunker import Chunk, chunk_text
from kb.vector_store import VectorStore

logger = logging.getLogger(__name__)

# Suppress noisy library warnings
warnings.filterwarnings("ignore", message=".*XMLParser.*", category=UserWarning)


# ── Text-based file readers ─────────────────────────────────────────

def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


_READERS = {
    ".txt": _read_text,
    ".md": _read_text,
    ".py": _read_text,
    ".json": _read_text,
    ".yaml": _read_text,
    ".yml": _read_text,
    ".csv": _read_text,
    ".html": _read_text,
    ".xml": _read_text,
    ".rst": _read_text,
}


# ── PDF reader ──────────────────────────────────────────────────────

def _read_pdf(path: Path) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF support. Install it with: pip install PyPDF2>=3.0.0")

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            if text:
                parts.append(text)
        except Exception as exc:
            logger.warning("[ingest] failed to extract text from PDF page %d of %s: %s", i, path, exc)

    if not parts:
        logger.warning("[ingest] no text extracted from PDF: %s", path)
        return ""

    return "\n\n".join(parts)


# ── DOCX reader ─────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX support. Install it with: pip install python-docx>=1.1.0")

    doc = Document(str(path))
    parts = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    if not parts:
        logger.warning("[ingest] no text extracted from DOCX: %s", path)
        return ""

    return "\n\n".join(parts)


# Add PDF and DOCX to the readers
_READERS[".pdf"] = _read_pdf
_READERS[".docx"] = _read_docx


class DocumentIngestionError(Exception):
    """Raised when a document cannot be ingested."""
    pass


def ingest_documents(
    store: VectorStore,
    data_dir: str | None = None,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> int:
    """Walk a directory, read all supported files, chunk, and add to the vector store.

    Args:
        store: The VectorStore to add chunks to.
        data_dir: Directory to scan. Defaults to config.kb.data_dir.
        chunk_size: Target chunk size. Defaults to config.kb.chunk_size.
        chunk_overlap: Overlap between chunks. Defaults to config.kb.chunk_overlap.

    Returns:
        The total number of chunks indexed.

    Raises:
        DocumentIngestionError: If a critical error occurs during ingestion.
    """
    data_dir = data_dir or config.kb.data_dir
    chunk_size = chunk_size or config.kb.chunk_size
    chunk_overlap = chunk_overlap or config.kb.chunk_overlap

    os.makedirs(data_dir, exist_ok=True)
    root = Path(data_dir)
    total_chunks = 0
    file_count = 0
    error_count = 0

    # Known directories to skip (session storage, etc.)
    _SKIP_DIRS = {"sessions", ".git", ".venv", "__pycache__", ".claude"}

    for file_path in sorted(root.rglob("*")):
        if file_path.is_dir():
            continue
        if any(part in _SKIP_DIRS for part in file_path.parts):
            continue

        ext = file_path.suffix.lower()
        reader = _READERS.get(ext)
        if reader is None:
            logger.debug("[ingest] skipping unsupported file: %s", file_path)
            continue

        try:
            text = reader(file_path)
        except ImportError as exc:
            logger.error("[ingest] missing dependency for %s: %s", file_path, exc)
            error_count += 1
            continue
        except Exception as exc:
            logger.error("[ingest] failed to read %s: %s", file_path, exc)
            error_count += 1
            continue

        if not text or not text.strip():
            logger.warning("[ingest] empty content from: %s", file_path)
            continue

        try:
            chunks = chunk_text(
                text,
                source=str(file_path),
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            store.add_chunks(chunks)
            total_chunks += len(chunks)
            file_count += 1
            logger.info("[ingest] indexed %s: %d chunks", file_path.name, len(chunks))
        except Exception as exc:
            logger.error("[ingest] failed to chunk/index %s: %s", file_path, exc)
            error_count += 1

    logger.info(
        "[ingest] completed: %d files, %d chunks, %d errors",
        file_count, total_chunks, error_count
    )

    if error_count > 0 and file_count == 0:
        raise DocumentIngestionError(
            f"All {error_count} file(s) failed to ingest. Check logs for details."
        )

    return total_chunks


def ingest_text(
    store: VectorStore,
    text: str,
    source: str = "api-upload",
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> int:
    """Ingest a single text string into the vector store.

    Args:
        store: The VectorStore to add chunks to.
        text: The text to ingest.
        source: Source identifier for the chunks.
        chunk_size: Target chunk size. Defaults to config.kb.chunk_size.
        chunk_overlap: Overlap between chunks. Defaults to config.kb.chunk_overlap.

    Returns:
        The number of chunks created.
    """
    if not text or not text.strip():
        logger.warning("[ingest] empty text provided, nothing to ingest")
        return 0

    chunk_size = chunk_size or config.kb.chunk_size
    chunk_overlap = chunk_overlap or config.kb.chunk_overlap

    chunks = chunk_text(text, source=source, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    store.add_chunks(chunks)
    return len(chunks)


def ingest_file(
    store: VectorStore,
    file_path: str | Path,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> int:
    """Ingest a single file into the vector store.

    Args:
        store: The VectorStore to add chunks to.
        file_path: Path to the file to ingest.
        chunk_size: Target chunk size. Defaults to config.kb.chunk_size.
        chunk_overlap: Overlap between chunks. Defaults to config.kb.chunk_overlap.

    Returns:
        The number of chunks created.

    Raises:
        DocumentIngestionError: If the file cannot be read or is unsupported.
    """
    path = Path(file_path)
    if not path.exists():
        raise DocumentIngestionError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise DocumentIngestionError(f"Unsupported file type: {ext}. Supported: {list(_READERS.keys())}")

    try:
        text = reader(path)
    except ImportError as exc:
        raise DocumentIngestionError(f"Missing dependency for {ext}: {exc}")
    except Exception as exc:
        raise DocumentIngestionError(f"Failed to read {file_path}: {exc}")

    if not text or not text.strip():
        raise DocumentIngestionError(f"Empty content from: {file_path}")

    chunk_size = chunk_size or config.kb.chunk_size
    chunk_overlap = chunk_overlap or config.kb.chunk_overlap

    chunks = chunk_text(text, source=str(path), chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    store.add_chunks(chunks)
    return len(chunks)
